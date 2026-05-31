# -*- coding: utf-8 -*-
"""Genera el documento Word profesional del procedimiento completo de SolUno."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paleta de marca SolUno
# ---------------------------------------------------------------------------
INDIGO = RGBColor(0x2A, 0x2A, 0x6E)      # confianza
NARANJA = RGBColor(0xE8, 0x7A, 0x16)     # sol BancoSol
GRIS = RGBColor(0x55, 0x55, 0x55)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
INDIGO_HEX = "2A2A6E"
NARANJA_HEX = "E87A16"
GRIS_CLARO_HEX = "EDEDF4"

doc = Document()

# ---------------------------------------------------------------------------
# Estilos base
# ---------------------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for hname, color, size in [("Heading 1", INDIGO, 16), ("Heading 2", NARANJA, 13), ("Heading 3", INDIGO, 12)]:
    st = doc.styles[hname]
    st.font.name = "Calibri"
    st.font.color.rgb = color
    st.font.size = Pt(size)
    st.font.bold = True


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_table(headers, rows, col_widths=None, header_fill=INDIGO_HEX):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=BLANCO, size=10)
        shade_cell(hdr[i], header_fill)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=10)
            if r_idx % 2 == 1:
                shade_cell(cells[i], GRIS_CLARO_HEX)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def bullets(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        if isinstance(it, tuple):
            run = p.add_run(it[0] + ": ")
            run.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def numbered(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        if isinstance(it, tuple):
            run = p.add_run(it[0] + ": ")
            run.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def para(text, bold=False, italic=False, align=None, color=None, size=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = GRIS
    # barra lateral
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), NARANJA_HEX)
    pbdr.append(left)
    pPr.append(pbdr)
    return p


# ===========================================================================
# 1. PORTADA
# ===========================================================================
for _ in range(4):
    doc.add_paragraph()

p = para("SolUno", align=WD_ALIGN_PARAGRAPH.CENTER, color=INDIGO)
p.runs[0].font.size = Pt(54)
p.runs[0].bold = True

p = para("Procedimiento Completo de la Plataforma", align=WD_ALIGN_PARAGRAPH.CENTER, color=NARANJA)
p.runs[0].font.size = Pt(22)
p.runs[0].bold = True

p = para("Un punto. Una respuesta. Todos los canales.", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
p.runs[0].font.size = Pt(14)
p.runs[0].italic = True

doc.add_paragraph()
p = para("Sistema operativo de incidencias y comunicacion para BancoSol",
         align=WD_ALIGN_PARAGRAPH.CENTER, color=INDIGO)
p.runs[0].font.size = Pt(13)

for _ in range(6):
    doc.add_paragraph()

p = para("InnovaHack 2026  ·  BancoSol", align=WD_ALIGN_PARAGRAPH.CENTER, color=INDIGO)
p.runs[0].font.size = Pt(13)
p.runs[0].bold = True
p = para("Equipo SolUno", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
p.runs[0].font.size = Pt(12)
p = para("30 de mayo de 2026", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRIS)
p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ===========================================================================
# TABLA DE CONTENIDO (manual, simple)
# ===========================================================================
h1("Contenido")
toc = [
    "1.  Resumen ejecutivo",
    "2.  El problema de BancoSol",
    "3.  Problem Statement validado por el banco",
    "4.  Arquitectura general (on-premise vs. cloud)",
    "5.  El procedimiento paso a paso",
    "6.  Tabla resumen del ruteo a equipos",
    "7.  Cumplimiento y privacidad",
    "8.  Stack tecnico",
    "9.  KPIs e impacto esperado",
    "10. Conclusion",
]
for t in toc:
    p = doc.add_paragraph(t)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ===========================================================================
# 1. RESUMEN EJECUTIVO
# ===========================================================================
h1("1. Resumen ejecutivo")
para(
    "SolUno es un punto unico que recibe cada queja e incidente desde todos los canales de BancoSol, "
    "reconoce de que se trata, a cuantos clientes afecta y con que datos reproducirlo, entrega una "
    "respuesta unica y correcta, rutea los incidentes al equipo tecnico de forma estructurada y formal "
    "(reemplazando los grupos de WhatsApp informales) y devuelve la solucion a un manual vivo que mantiene "
    "a call center y clientes hablando el mismo idioma, todo sin que un dato confidencial salga del banco."
)
quote("En una frase: Un punto. Una respuesta. Todos los canales.")
para(
    "Este documento describe el procedimiento completo de la plataforma paso a paso, con foco en el flujo "
    "web/tecnico: como ingresan los reportes, como se normalizan y tokenizan, como se clasifican, la "
    "bifurcacion clave entre queja e incidente, el ruteo a los equipos del banco, la agrupacion de "
    "reportes, el cierre del ciclo hacia el Playbook Vivo y la respuesta al cliente."
)

# ===========================================================================
# 2. EL PROBLEMA
# ===========================================================================
h1("2. El problema de BancoSol")
para(
    "BancoSol es la institucion pionera y lider mundial en microfinanzas de Bolivia: mas de 2,5 millones de "
    "clientes, +1.801 puntos de atencion, #1 del ranking CAMEL durante 10 anos consecutivos y una cartera de "
    "USD 2.712 millones, bajo supervision estricta de ASFI. En mayo de 2024 lanzo altoke (billetera QR para "
    "microempresarios), que paso de 0 a 1,2 millones de usuarios y disparo el trafico de API en +2.000% en "
    "menos de 6 meses, multiplicando incidencias y reclamos en todos los canales."
)
h2("2.1 El dolor real")
bullets([
    ("Canales fragmentados", "los reclamos llegan por 9 vias distintas (WhatsApp, FonoSol, appSol, Solnet, altoke, sucursales, chatbot, redes sociales, Punto de Reclamo web), cada una con su propio playbook y ritmo de actualizacion."),
    ("Respuestas inconsistentes", "un mismo problema recibe respuestas distintas segun el canal."),
    ("Call Center cuello de botella", "tercerizado, sin conocimiento tecnico profundo; deriva el resto a desarrollo en lotes diarios, por lo que los devs hacen soporte en vez de desarrollar."),
    ("Comunicacion interna informal", "los incidentes se reportan en grupos de Teams/WhatsApp informales; el equipo tecnico no sabe a cuantos usuarios afecta ni recibe datos para replicar el error, y al resolver la solucion no vuelve ordenada al playbook."),
    ("El cliente no tiene a donde ir", "no existe un punto unico para reportar, consultar estado o auto-resolver."),
])
h2("2.2 La distincion que dio el banco: Queja vs. Incidente")
add_table(
    ["", "QUEJA / CONSULTA", "INCIDENTE"],
    [
        ["Que es", "Problema particular, bajo impacto, error controlado y conocido", "Afectacion grave: falla una funcion critica"],
        ["Ejemplos", "Extranjero no transacciona tras update (0,5%); cambio de dispositivo a reactivar token; 'pague y sigue la deuda'", "Falla login, transferencia, QR, ACH"],
        ["Como se resuelve", "Operativo y autonomo: playbook actualizado, misma respuesta en todo canal", "Se compila y prioriza informacion y se rutea al equipo tecnico"],
    ],
    col_widths=[1.3, 2.6, 2.6],
    header_fill=NARANJA_HEX,
)

# ===========================================================================
# 3. PROBLEM STATEMENT
# ===========================================================================
h1("3. Problem Statement validado por el banco")
para("Cita textual del Problem Statement final, validado por la persona del banco en la sesion de validacion:", italic=True)
quote(
    "El equipo de atencion de BancoSol -call center y equipo de desarrollo- necesita recibir en un unico "
    "punto y de forma estructurada cada queja e incidente que hoy llega disperso por todos los canales, "
    "sabiendo de que se trata, a cuantos clientes afecta y con los datos necesarios para reproducirlo, para "
    "entregar una respuesta unica, correcta y oportuna y, una vez resuelto, devolver esa solucion a un manual "
    "vivo que mantenga a todos -call center y cliente- hablando el mismo idioma."
)
quote(
    "Hoy no lo logran porque no existe ni siquiera un canal interno formal: los incidentes se reportan en "
    "grupos de chat informales, el equipo tecnico no sabe cuantos usuarios estan afectados ni recibe datos "
    "para replicar el error, y las soluciones no vuelven de forma ordenada a los manuales, que quedan "
    "desactualizados y distintos en cada canal."
)
quote(
    "Una solucion comercial existente no lo resuelve porque no puede operar sin que la informacion del "
    "cliente salga del banco, no entiende la logica especifica de cada canal y producto, y no mantiene un "
    "unico manual vivo y consistente que cumpla las exigencias de privacidad y de reporte individual del "
    "regulador boliviano."
)
para("Estado: Validado por la persona del banco ('Esta bien, esta bien').", bold=True, color=INDIGO)

# ===========================================================================
# 4. ARQUITECTURA
# ===========================================================================
h1("4. Arquitectura general")
para(
    "SolUno se apoya en una separacion estricta entre lo que corre dentro del banco (on-premise) y lo que "
    "corre en la nube. El principio rector es que la informacion identificable del cliente nunca sale del "
    "banco: el modelo de IA en la nube solo recibe texto ya tokenizado."
)
h2("4.1 Que corre donde")
add_table(
    ["Capa", "Componentes", "Datos que maneja"],
    [
        ["On-premise (dentro del banco)", "n8n (orquestador), vault de identidad, modulo de tokenizacion / re-hidratacion, base de playbooks versionada, base de incidentes, registro auditable", "PII real: nombre, CI, cuenta, telefono, saldos"],
        ["Cloud (Zero Data Retention)", "El LLM (Claude) que clasifica, redacta respuestas y genera informes", "Solo texto tokenizado: ve 'usuario 1', nunca 'Fernando Alen'"],
    ],
    col_widths=[1.8, 3.0, 2.0],
)
h2("4.2 Los tres pilares tecnicos")
bullets([
    ("n8n (orquestador)", "recibe los canales, normaliza, tokeniza, llama al cerebro de IA, bifurca el flujo y dispara las respuestas. Es el director de orquesta on-premise."),
    ("Claude (cerebro)", "clasifica queja vs. incidente, determina producto/categoria/severidad, redacta la respuesta unica y genera el informe ASFI, operando siempre sobre datos tokenizados."),
    ("Tokenizacion", "capa de privacidad que reemplaza la PII por tokens antes de cualquier IA y la re-hidrata localmente dentro del banco."),
])

# ===========================================================================
# 5. PROCEDIMIENTO PASO A PASO
# ===========================================================================
h1("5. El procedimiento paso a paso")
para(
    "Esta es la seccion central del documento. Describe, en orden, el recorrido completo de un reporte desde "
    "que entra por cualquier canal hasta que se cierra el ciclo y el cliente recibe su respuesta.",
    italic=True,
)

# 5.a Ingesta
h2("5.a Ingesta: recepcion desde WhatsApp y los canales Meta")
para(
    "El MVP recibe reportes por dos mundos: altoke (canal interno/transaccional, via su webhook/API interna) "
    "y Meta (canales externos y publicos). Dentro de Meta conviven tres canales:"
)
bullets([
    ("WhatsApp", "ingresa por el nodo nativo de WhatsApp en n8n (WhatsApp Cloud API)."),
    ("Facebook Messenger e Instagram", "ingresan por un webhook de Meta. Un mismo webhook de Meta Business recibe tanto Messenger como Instagram (Graph API), de modo que un solo punto de entrada cubre ambos canales publicos."),
])
para(
    "En la practica conviven dos puntos de ingesta: el nodo nativo de WhatsApp y el webhook de Meta que "
    "atiende Messenger e Instagram. El webhook generico de SolUno (POST a /webhook/soluno-reporte) recibe el "
    "evento con la estructura {canal, producto, mensaje, cliente:{nombre, ci, cuenta, telefono}}.",
)

# 5.b Normalizacion
h2("5.b Normalizacion: un formato comun")
para(
    "Cada canal trae su propia estructura (WhatsApp, Messenger e Instagram tienen payloads distintos; altoke "
    "agrega contexto transaccional). El primer trabajo de n8n es convertir cada evento a un formato comun "
    "unico, con los mismos campos (canal de origen, producto, texto del problema y bloque de cliente). A "
    "partir de aqui el resto del flujo es identico, sin importar por donde haya entrado el reporte."
)

# 5.c Tokenizacion
h2("5.c Tokenizacion (privacidad)")
para(
    "Antes de que cualquier IA toque el contenido, el modulo de tokenizacion on-premise reemplaza la PII por "
    "tokens. La tokenizacion es inteligente segun el tipo de dato: parte se conserva visible solo para el "
    "revisor interno y nada identificable viaja al LLM."
)
add_table(
    ["Dato", "Accion", "Por que"],
    [
        ["Nombre completo", "Tokenizar al LLM, visible al revisor", "El revisor lo necesita; el LLM no"],
        ["CI / Carnet", "Tokenizar al LLM, visible al revisor", "Suficiente para identificar internamente"],
        ["Nº de cuenta", "Tokenizar y ocultar", "Irrelevante para resolver"],
        ["Saldo / montos", "Tokenizar y ocultar", "Sensible, no necesario"],
        ["Telefono / correo / direccion", "Tokenizar y ocultar", "Irrelevante para la revision"],
        ["Texto del problema", "Pasa tokenizado al LLM", "Lo necesario para clasificar"],
    ],
    col_widths=[1.9, 2.4, 2.5],
)
para(
    "Se conserva nombre + CI visibles unicamente para el revisor interno; el resto (cuenta, saldos, telefono) "
    "se oculta. El LLM externo opera sobre 'usuario 1' y la re-hidratacion se hace localmente dentro del "
    "banco. Como el modelo nunca recibe datos identificables, no hay cesion a terceros."
)

# 5.d Clasificacion
h2("5.d Clasificacion")
para(
    "Con el texto ya tokenizado, SolUno clasifica el reporte en varias dimensiones a la vez:"
)
bullets([
    ("Tipo", "es una QUEJA o un INCIDENTE."),
    ("Producto", "appSol, altoke, Solnet, etc. (cada uno tiene su logica)."),
    ("Categoria", "acceso/login, pagos/QR, transferencias, rendimiento, etc."),
    ("Severidad", "para priorizar el caso si es incidente."),
])
para(
    "En el MVP la clasificacion arranca por reglas (palabras clave) y en el roadmap se reemplaza por Claude "
    "para mayor precision. El resultado de esta clasificacion es lo que decide la bifurcacion del siguiente paso."
)

# 5.e BIFURCACION
h2("5.e Bifurcacion clave: queja vs. incidente")
para(
    "Este es el corazon del procedimiento. Segun la clasificacion, el flujo se separa en dos caminos "
    "completamente distintos:",
    bold=True,
)
h3 = doc.add_heading("Camino QUEJA (resolucion operativa y autonoma)", level=3)
para(
    "Si es una queja (bajo impacto, error conocido), SolUno la resuelve de forma operativa y autonoma con el "
    "Playbook Vivo. Sirve la respuesta estandarizada y vigente, de modo que el cliente recibe la misma "
    "respuesta correcta en todos los canales y puede resolverlo el mismo (autoservicio). No se pasa a ningun "
    "equipo: el caso se cierra en este punto."
)
doc.add_heading("Camino INCIDENTE (ruteo al equipo correcto del banco)", level=3)
para(
    "Si es un incidente (falla grave de una funcion critica: login, transferencia, QR, ACH), SolUno arma un "
    "Paquete de Incidente y lo rutea al equipo correcto del banco. El Paquete de Incidente contiene:"
)
bullets([
    "Que falla (descripcion del problema).",
    "A cuantos clientes afecta (conteo de afectados).",
    "Datos para reproducirlo (tokenizados, con contexto transaccional cuando viene de altoke).",
    "Severidad / prioridad.",
    "Solo lo necesario: no toda la informacion se pasa al equipo tecnico.",
])
para(
    "Este Paquete de Incidente reemplaza los grupos de WhatsApp informales por un canal interno formal hacia "
    "el equipo responsable. El ruteo depende del tipo de incidente (ver tabla completa en la seccion 6):"
)
bullets([
    ("Pagos/QR (altoke, transacciones)", "Equipo Pagos/QR."),
    ("Acceso/login", "Equipo Seguridad/Accesos."),
    ("Apps caidas (appSol/Solnet)", "DevOps Aplicaciones."),
    ("Rendimiento/infraestructura", "Infraestructura/TI."),
    ("Otros", "Mesa de Incidentes TI."),
])

# 5.f Agrupacion
h2("5.f Agrupacion de reportes")
para(
    "Muchos reportes del mismo incidente se agrupan en UN solo incidente con conteo de afectados. Si 200 "
    "clientes reportan 'el QR no funciona' por distintos canales, SolUno no genera 200 incidentes sino 1 "
    "incidente con 200 afectados ('200 reclamos a 1 incidente'). Asi el equipo tecnico ve el problema real y "
    "su magnitud, en lugar de cientos de tickets sueltos."
)

# 5.g Cierre del ciclo
h2("5.g Cierre del ciclo: retorno al Playbook Vivo")
para(
    "Cuando el equipo resuelve el incidente, la solucion no se queda en un chat. SolUno la convierte en pasos "
    "1-2-3-4 que cualquiera pueda seguir y la publica al Playbook Vivo, que es versionado y trazable. De esta "
    "forma, el proximo caso identico ya tiene respuesta lista y se propaga a todos los canales al instante. "
    "Esto cierra el ciclo: una queja futura del mismo tipo se resuelve sola por el camino operativo."
)

# 5.h Respuesta al cliente
h2("5.h Respuesta al cliente")
para(
    "La respuesta se envia al cliente por el mismo canal por el que reporto (WhatsApp, Messenger, Instagram o "
    "altoke). El cliente sigue el estado de su caso en tres estados claros:"
)
add_table(
    ["Estado", "Significado"],
    [
        ["Recibido", "SolUno recibio el reporte y lo registro."],
        ["En analisis", "El caso esta clasificado y, si es incidente, en manos del equipo correcto."],
        ["Resuelto", "Se aplico la solucion del playbook o el equipo cerro el incidente."],
    ],
    col_widths=[2.0, 4.5],
    header_fill=NARANJA_HEX,
)

# 5.i Registro y metricas
h2("5.i Registro de estado y metricas")
para(
    "Cada paso queda registrado en la base de incidentes y en el registro auditable on-premise. A partir de "
    "ese registro, SolUno produce metricas operativas:"
)
bullets([
    "Tiempos (de recepcion a resolucion).",
    "Casos por equipo (carga de cada area).",
    "Recurrencia (problemas que se repiten y deberian ir al playbook).",
    "Picos (momentos criticos de volumen, util para detectar incidentes masivos temprano).",
])

# 5.j Plus
h2("5.j Plus: informe ASFI y traduccion de idiomas")
para(
    "SolUno suma dos capacidades extra que multiplican su valor:",
)
bullets([
    ("Informe automatico para ASFI (Plus #1)", "al cerrar un incidente, el LLM redacta el informe individualizado/post-mortem (que paso, en que servidor, como se resolvio, que acciones evitan que se repita). El humano revisa y aprueba. Cumple el reporte regulatorio y alimenta el playbook."),
    ("Traduccion de idiomas (Plus #2)", "traduce el reporte del cliente (quechua, aymara, guarani o un idioma extranjero) al espanol para el operador, y la respuesta del playbook al idioma del cliente. Encaja con la mision social e inclusiva de BancoSol."),
])

# ===========================================================================
# 6. TABLA DE RUTEO
# ===========================================================================
h1("6. Tabla resumen del ruteo a equipos")
para(
    "Cuando un reporte se clasifica como incidente, el Paquete de Incidente se rutea al equipo responsable "
    "segun el tipo. Esta tabla resume la logica de ruteo:"
)
add_table(
    ["Tipo de incidente", "Ejemplos", "Equipo responsable"],
    [
        ["Pagos / QR", "altoke, transacciones por QR, cobros/pagos", "Equipo Pagos/QR"],
        ["Acceso / login", "no puede iniciar sesion, reactivacion de token", "Equipo Seguridad/Accesos"],
        ["Apps caidas", "appSol o Solnet no abren / caidas", "DevOps Aplicaciones"],
        ["Rendimiento / infraestructura", "lentitud, timeouts, saturacion de servidores", "Infraestructura/TI"],
        ["Otros", "casos no categorizados", "Mesa de Incidentes TI"],
    ],
    col_widths=[2.0, 2.7, 2.0],
)
para(
    "Una queja, en cambio, no se rutea a ningun equipo: se resuelve de forma operativa con el Playbook Vivo y "
    "se responde directamente al cliente.",
    italic=True,
)

# ===========================================================================
# 7. CUMPLIMIENTO Y PRIVACIDAD
# ===========================================================================
h1("7. Cumplimiento y privacidad")
para(
    "La premisa del banco es clara: 'Toda informacion del cliente tiene que quedarse en el banco.' SolUno la "
    "respeta por diseno:"
)
bullets([
    "La PII se tokeniza on-premise antes de cualquier llamada a la IA.",
    "El LLM en la nube opera sobre 'usuario 1' y nunca recibe datos identificables (Zero Data Retention).",
    "La re-hidratacion (volver del token al dato real) ocurre localmente dentro del banco.",
    "Como el modelo externo no recibe datos identificables, no hay cesion a terceros, lo que simplifica el cumplimiento de privacidad y de tercerizacion ante ASFI.",
    "Cada reporte genera un registro auditable, alineado con el requisito de expediente individual del regulador.",
])
para(
    "El marco regulatorio (ASFI: respuesta en 5 dias habiles, Punto de Reclamo, expediente individual, "
    "soberania del dato) es una restriccion a respetar, no el problema que SolUno resuelve. El core es la "
    "comunicacion; el informe ASFI automatico es un plus que ayuda a cumplir el reporte regulatorio."
)

# ===========================================================================
# 8. STACK TECNICO
# ===========================================================================
h1("8. Stack tecnico")
add_table(
    ["Componente", "Rol", "Detalle"],
    [
        ["n8n", "Orquestador (on-premise)", "Webhook de ingesta, normalizacion, tokenizacion, clasificacion por reglas, bifurcacion queja/incidente, armado del Paquete de Incidente, respuesta. Nodos nativos de WhatsApp y webhook de Meta para Messenger/Instagram."],
        ["Claude (Anthropic)", "Cerebro de IA (cloud, ZDR)", "Clasifica, redacta la respuesta unica y genera el informe ASFI sobre datos tokenizados. La API key se guarda dentro de n8n (Credentials), nunca en el repo."],
        ["Frontend", "Vistas de usuario", "Dos vistas (cliente con 3 estados; interna: call center, desarrollo y cierre del ciclo) + status + metricas, consumiendo los webhooks (React / v0 / Lovable)."],
        ["Base de datos", "Persistencia on-premise", "Base de playbooks versionada, base de incidentes (agrupacion y conteo de afectados, Postgres/pgvector) y registro auditable."],
    ],
    col_widths=[1.5, 1.8, 3.4],
)
h2("8.1 Flujo de nodos en n8n (workflow v1)")
numbered([
    ("Webhook (Ingesta)", "recibe POST a /webhook/soluno-reporte."),
    ("Tokenizar (PII)", "reemplaza nombre/CI/cuenta/telefono por tokens; conserva nombre+CI para el revisor."),
    ("Clasificar (reglas)", "queja vs. incidente, producto, categoria, severidad."),
    ("¿Es Incidente?", "bifurca: incidente al paso 5a; queja al paso 5b."),
    ("Paquete de Incidente (5a)", "qué falla, # afectados, datos para reproducir, equipo, estado."),
    ("Respuesta Queja / Playbook (5b)", "trae la misma respuesta y pasos del playbook segun categoria."),
    ("Responder", "devuelve el JSON final y dispara la respuesta al canal de origen."),
])

# ===========================================================================
# 9. KPIs E IMPACTO
# ===========================================================================
h1("9. KPIs e impacto esperado")
para(
    "Los siguientes datos son publicos y verificables; el banco no entrega cifras reales, por lo que se usan "
    "para construir supuestos creibles. Cada fila indica su fuente.",
    italic=True,
)
add_table(
    ["Indicador / dato", "Cifra", "Fuente"],
    [
        ["Escala de BancoSol", "+2,5M clientes, #1 CAMEL 10 anos", "bancosol.com.bo"],
        ["Crecimiento de altoke", "0 a 1,2M usuarios; +2.000% de API en 6 meses", "Sensedia"],
        ["Llamadas repetidas al call center", "hasta 32%", "Balto / Enghouse"],
        ["Tiempo de dev en soporte/mantenimiento", "17,3 h de 41,1 por semana", "Stripe Developer Coefficient"],
        ["Costo de 1 hora de caida", "+USD 300.000 (90% de empresas)", "ITIC 2024"],
        ["Plazo legal ASFI de respuesta", "5 dias habiles", "Reglamento ASFI"],
        ["Reclamos 1a instancia 2024", "75.805 (a verificar)", "ASFI CIRS"],
        ["Esperan misma respuesta en todo canal", "75%", "Zendesk CX Trends 2024"],
        ["Deben repetir su info a distintos agentes", "66%", "Salesforce State of Service"],
        ["Ahorro de GenAI en atencion", "30-45% de los costos de la funcion", "McKinsey"],
        ["Hablantes de lenguas originarias", "2,2M (quechua 1,39M, aymara 774.874, guarani 43.870)", "Censo 2024 (INE)"],
    ],
    col_widths=[2.7, 2.2, 1.8],
)
h2("9.1 Impacto esperado de SolUno")
bullets([
    ("Menos respuestas inconsistentes", "una sola respuesta correcta por canal gracias al Playbook Vivo."),
    ("Devs liberados", "los incidentes llegan estructurados y agrupados, no como tickets sueltos; el equipo deja de hacer soporte manual."),
    ("Resolucion mas rapida de incidentes masivos", "la agrupacion y el conteo de afectados muestran la magnitud real temprano."),
    ("Cumplimiento mas simple", "PII que no sale del banco e informe ASFI automatico."),
    ("Inclusion", "atencion en lenguas originarias y a clientes extranjeros."),
])

# ===========================================================================
# 10. CONCLUSION
# ===========================================================================
h1("10. Conclusion")
para(
    "SolUno convierte un proceso hoy disperso e informal en un procedimiento unico, estructurado y trazable. "
    "Recibe cada reporte por cualquier canal, protege la informacion del cliente con tokenizacion on-premise, "
    "clasifica y bifurca entre quejas (autoservicio operativo) e incidentes (ruteo formal al equipo correcto), "
    "agrupa los reportes en incidentes reales con su conteo de afectados, cierra el ciclo devolviendo cada "
    "solucion al Playbook Vivo y responde al cliente por su mismo canal con estados claros."
)
para(
    "El resultado es un banco que deja de improvisar la comunicacion: call center y clientes hablan el mismo "
    "idioma, los desarrolladores vuelven a desarrollar y el regulador recibe lo que exige, todo sin que un "
    "dato confidencial salga del banco.",
)
quote("SolUno: un punto, una respuesta, todos los canales.")

# ---------------------------------------------------------------------------
out = r"c:\Users\ASUS\3D Objects\HACKATON\SolUno_Procedimiento_Completo.docx"
doc.save(out)
print("OK ->", out)
